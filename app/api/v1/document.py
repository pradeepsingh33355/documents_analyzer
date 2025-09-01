import os
import uuid
from fastapi import FastAPI, UploadFile, File, HTTPException, APIRouter, Depends, status
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
from typing import Optional, Dict
from app.agent.agent import GrammarCorrectionAgent
import shutil
from fastapi import FastAPI, Response
from fastapi.requests import Request
from pathlib import Path
import os
from docx import Document as DocxDocument
from app.core.config import settings
from app.schemas.schemas import User
from app.api.v1.auth import get_current_user
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.models.user import Document
from app.services.correct_file import correct_text
from app.utils.extract import guess_and_extract
from app.ai.compliance import assess_text
from app.schemas.CorrectionRequest import CorrectionRequest

router = APIRouter()

DOC_STORE: Dict[str, Dict] = {}

UPLOAD_DIR = settings.UPLOAD_DIR


UPLOAD_DIR = "app/storage/original_file"
os.makedirs(UPLOAD_DIR, exist_ok=True)


@router.post("/upload")
async def upload(
    file: UploadFile,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    name = file.filename

    if not (name.lower().endswith(".pdf") or name.lower().endswith(".docx")):
        raise HTTPException(
            status_code=400, detail="Only .pdf and .docx files are allowed."
        )

    user_id = current_user.id
    filename_with_user = (
        f"{os.path.splitext(name)[0]}_{user_id}{os.path.splitext(name)[1]}"
    )

    file_location = os.path.join(UPLOAD_DIR, filename_with_user)

    with open(file_location, "wb") as f:
        f.write(await file.read())

    document = Document(
        original_file=filename_with_user, status="pending", owner_id=user_id
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    return JSONResponse(
        content={
            "status": "uploaded",
            "saved_as": filename_with_user,
            "path": file_location,
        },
        status_code=status.HTTP_200_OK,
    )


@router.post("/assess/{doc_id}")
async def assess(
    doc_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc_name = (
        db.query(Document.original_file)
        .filter(Document.id == doc_id)
        .first()
        .original_file
    )

    file_path = os.path.join(UPLOAD_DIR, doc_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Document not found.")
    with open(file_path, "rb") as f:
        data = f.read()

    text = guess_and_extract(doc_name, data)

    report = assess_text(text)
    return JSONResponse({"doc_id": doc_id, **report}, status_code=status.HTTP_200_OK)


@router.post("/correct_all_errors")
async def correct_all_errors(
    request: CorrectionRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_id = request.doc_id
    doc_name = (
        db.query(Document.original_file)
        .filter(Document.id == doc_id)
        .first()
        .original_file
    )

    base_doc_name, extension = os.path.splitext(doc_name)
    user_id = current_user.id
    filename_with_user = (
        f"{os.path.splitext(doc_name)[0]}_{'corrected'}{os.path.splitext(doc_name)[1]}"
    )

    corrected_file_path = str(
        Path("app/storage/corrected_files") / f"{base_doc_name}_corrected.docx"
    )

    if extension not in [".pdf", ".docx"]:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    file_path = os.path.join(UPLOAD_DIR, doc_name)
    try:
        with open(file_path, "rb") as f:
            data = f.read()
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="File not found")

    text = guess_and_extract(doc_name, data)

    issues = request
    if not issues:
        raise HTTPException(status_code=400, detail="No issues to correct")

    agent = GrammarCorrectionAgent(doc_id=doc_id)

    guidelines = request.guideline_notes
    corrected_text = await agent.run(issues=issues, guidelines=guidelines, text=text)

    doc = DocxDocument()
    doc.add_paragraph(corrected_text)
    doc.save(corrected_file_path)

    existing_document = db.query(Document).filter(Document.id == doc_id).first()
    if existing_document:
        existing_document.corrected_file = filename_with_user
        existing_document.status = "Completed"
        db.commit()
        db.refresh(existing_document)
    else:
        raise HTTPException(status_code=404, detail="Document not found")

    return JSONResponse(
        content={"corrected_text": corrected_text},
        status_code=status.HTTP_200_OK,
    )


@router.post("/download/{doc_id}")
async def create_and_download_file(
    doc_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = db.query(Document).filter(Document.id == doc_id).first()

    original_file_name = document.original_file
    base_doc_name, extension = os.path.splitext(original_file_name)
    file_name = f"{base_doc_name}_corrected.docx"

    src_file_path = Path("app/storage/corrected_files") / file_name

    if not src_file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return {"message": f"File {file_name} downlaoded successfully"}
