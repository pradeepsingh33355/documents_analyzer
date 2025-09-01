import io
from fastapi.testclient import TestClient
from app.main import app
from app.utils.DemoJsonForCorrectApi import sample_json

client = TestClient(app)


def create_sample_docx():
    from docx import Document

    buf = io.BytesIO()
    d = Document()
    d.add_paragraph(
        "This is a smple documment. It have bad grammar and long sentence which is not so good because the clarity is low and its very very very long and complicated for no reason at all"
    )
    d.add_paragraph("Second paragraph with an eror.")
    d.save(buf)
    buf.seek(0)
    return buf


def test_upload_assess_correct_flow():
    # Auth(Login)
    login_data = {"email": "user@example.com", "password": "password"}
    r = client.post("/login", json=login_data)
    assert r.status_code == 200
    access_token = r.json()["access_token"]

    # Upload (File is uploaded in the provided path locally)
    buf = create_sample_docx()
    files = {
        "file": (
            "sample.docx",
            buf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    headers = {"Authorization": f"Bearer {access_token}"}
    r = client.post("/upload", files=files, headers=headers)
    assert r.status_code == 200
    doc_id = r.json()["doc_id"]

    # Assess(file is analyzed)
    r2 = client.post(f"/assess/{doc_id}", headers=headers)
    assert r2.status_code == 200
    data = r2.json()
    assert "scores" in data
    assert "grammar_issues" in data["scores"]

    # Correct(File is corrected and downloaded)
    r3 = client.post(
        f"/correct_all_errors/{doc_id}",
        json=sample_json,
        headers=headers,
    )
    assert r3.status_code == 200
    assert r3.headers.get("content-disposition").startswith("attachment;")
